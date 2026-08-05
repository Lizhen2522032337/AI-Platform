"""报表生成与文件 Tool：Markdown、Word、PDF、Excel、JSON 和 CSV。"""

import csv
import html
import io
import json
import logging
import re
from dataclasses import dataclass
from typing import Any, Literal

from app.config.settings import Settings, get_settings
from app.integrations import save_file

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class MarkdownBlock:
    """供 Word、PDF 和 Excel 共同使用的轻量 Markdown 块。"""

    kind: Literal["heading", "paragraph", "bullet", "number", "quote", "table"]
    text: str = ""
    level: int = 0
    rows: list[list[str]] | None = None


def _safe_name(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "-", value).strip("-")
    return cleaned[:80] or "data"


def _plain_text(value: object) -> str:
    """移除常见 Markdown 标记，保留适合办公文档显示的文本。"""

    text = str(value or "")
    text = re.sub(r"\[([^]]+)]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    return text.replace("`", "").strip()


def _table_cells(line: str) -> list[str]:
    return [_plain_text(cell) for cell in line.strip().strip("|").split("|")]


def _is_table_separator(line: str) -> bool:
    cells = line.strip().strip("|").split("|")
    return bool(cells) and all(re.fullmatch(r"\s*:?-{3,}:?\s*", cell) for cell in cells)


def _parse_markdown(markdown: str) -> list[MarkdownBlock]:
    """解析平台当前会产生的标题、列表、引用和表格子集。"""

    lines = markdown.replace("\r\n", "\n").split("\n")
    blocks: list[MarkdownBlock] = []
    index = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(
                MarkdownBlock(kind="paragraph", text=_plain_text(" ".join(paragraph_lines)))
            )
            paragraph_lines.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if (
            "|" in stripped
            and index + 1 < len(lines)
            and _is_table_separator(lines[index + 1])
        ):
            flush_paragraph()
            rows = [_table_cells(stripped)]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                rows.append(_table_cells(lines[index]))
                index += 1
            column_count = max(len(row) for row in rows)
            normalized = [row + [""] * (column_count - len(row)) for row in rows]
            blocks.append(MarkdownBlock(kind="table", rows=normalized))
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    kind="heading",
                    text=_plain_text(heading.group(2)),
                    level=min(len(heading.group(1)), 3),
                )
            )
        elif re.match(r"^[-*+]\s+", stripped):
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    kind="bullet",
                    text=_plain_text(re.sub(r"^[-*+]\s+", "", stripped)),
                )
            )
        elif re.match(r"^\d+[.)]\s+", stripped):
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    kind="number",
                    text=_plain_text(re.sub(r"^\d+[.)]\s+", "", stripped)),
                )
            )
        elif stripped.startswith(">"):
            flush_paragraph()
            blocks.append(MarkdownBlock(kind="quote", text=_plain_text(stripped[1:])))
        elif stripped in {"---", "***", "___"}:
            flush_paragraph()
        else:
            paragraph_lines.append(stripped)
        index += 1
    flush_paragraph()
    return blocks


def _data_observations(
    observations: list[dict[str, Any]],
) -> list[tuple[str, list[str], list[dict[str, object]]]]:
    """提取所有成功的数据库结果，包括管理员平台动态查询。"""

    results: list[tuple[str, list[str], list[dict[str, object]]]] = []
    used_names: set[str] = set()
    for index, observation in enumerate(observations, start=1):
        if (
            observation.get("tool")
            not in {"database_query", "db2_query", "dynamic_sql"}
            or observation.get("status") != "ok"
        ):
            continue
        rows = observation.get("rows")
        columns = observation.get("columns")
        if not isinstance(rows, list) or not isinstance(columns, list):
            continue
        base_name = _safe_name(
            str(
                observation.get("queryId")
                or ("platform_data" if observation.get("tool") == "dynamic_sql" else f"query_{index}")
            )
        )
        name = base_name
        suffix = 2
        while name in used_names:
            name = f"{base_name}_{suffix}"
            suffix += 1
        used_names.add(name)
        normalized_rows = [row for row in rows if isinstance(row, dict)]
        results.append((name, [str(column) for column in columns], normalized_rows))
    return results


def _csv_payload(rows: list[dict[str, object]], columns: list[str]) -> bytes:
    stream = io.StringIO(newline="")
    writer = csv.DictWriter(stream, fieldnames=columns, extrasaction="ignore")
    writer.writeheader()
    for row in rows:
        writer.writerow(row)
    # Excel 在 Windows 上能自动识别 UTF-8 BOM。
    return stream.getvalue().encode("utf-8-sig")


def _set_docx_font(run: Any, ascii_name: str = "Calibri", east_asia: str = "微软雅黑") -> None:
    """同时设置西文字体和东亚字体，避免中文被错误替换。"""

    from docx.oxml.ns import qn

    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _docx_payload(title: str, markdown: str) -> bytes:
    """按标准业务报告样式生成可编辑 Word 文档。"""

    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    blocks = _parse_markdown(markdown)
    max_table_columns = max(
        (len(block.rows[0]) for block in blocks if block.kind == "table" and block.rows),
        default=0,
    )
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    if max_table_columns > 6:
        # 宽数据表采用横向页面，避免大量字段被挤压或截断。
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_after = Pt(14)
    title_run = title_paragraph.add_run(_plain_text(title))
    _set_docx_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor.from_string("0B2545")

    for block in blocks:
        if block.kind == "heading":
            paragraph = document.add_heading(block.text, level=max(1, block.level))
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(block.text, style="List Bullet")
        elif block.kind == "number":
            paragraph = document.add_paragraph(block.text, style="List Number")
        elif block.kind == "quote":
            paragraph = document.add_paragraph(block.text)
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.right_indent = Inches(0.2)
            for run in paragraph.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor.from_string("555555")
        elif block.kind == "table" and block.rows:
            table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            table.style = "Table Grid"
            table.autofit = True
            for row_index, source_row in enumerate(block.rows):
                for column_index, value in enumerate(source_row):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = value
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            _set_docx_font(run)
                            run.font.size = Pt(8 if len(source_row) > 6 else 9)
                            run.bold = row_index == 0
                    if row_index == 0:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "F2F4F7")
                        cell._tc.get_or_add_tcPr().append(shading)
            document.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        else:
            paragraph = document.add_paragraph(block.text)
        for run in paragraph.runs:
            _set_docx_font(run)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _pdf_inline(value: str) -> str:
    escaped = html.escape(value)
    escaped = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", escaped)
    return escaped.replace("`", "")


def _pdf_payload(title: str, markdown: str) -> bytes:
    """使用内置中文 CID 字体生成无需额外字体文件的 PDF。"""

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    blocks = _parse_markdown(markdown)
    max_table_columns = max(
        (len(block.rows[0]) for block in blocks if block.kind == "table" and block.rows),
        default=0,
    )
    page_size = landscape(letter) if max_table_columns > 6 else letter
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    stream = io.BytesIO()
    document = SimpleDocTemplate(
        stream,
        pagesize=page_size,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
        title=_plain_text(title),
        author="Enterprise AI Platform",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10,
        leading=16,
        textColor=colors.HexColor("#202124"),
        spaceAfter=7,
    )
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=body,
        fontSize=19,
        leading=25,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=16,
    )
    heading_styles = {
        1: ParagraphStyle(
            "ChineseHeading1",
            parent=body,
            fontSize=15,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=8,
        ),
        2: ParagraphStyle(
            "ChineseHeading2",
            parent=body,
            fontSize=12.5,
            leading=17,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=9,
            spaceAfter=6,
        ),
        3: ParagraphStyle(
            "ChineseHeading3",
            parent=body,
            fontSize=11,
            leading=16,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=7,
            spaceAfter=5,
        ),
    }
    story: list[Any] = [Paragraph(html.escape(_plain_text(title)), title_style)]
    for block in blocks:
        if block.kind == "heading":
            story.append(
                Paragraph(html.escape(block.text), heading_styles[max(1, block.level)])
            )
        elif block.kind in {"bullet", "number"}:
            marker = "•" if block.kind == "bullet" else "-"
            story.append(Paragraph(f"{marker} {_pdf_inline(block.text)}", body))
        elif block.kind == "quote":
            quote_style = ParagraphStyle(
                "ChineseQuote",
                parent=body,
                leftIndent=16,
                textColor=colors.HexColor("#555555"),
            )
            story.append(Paragraph(_pdf_inline(block.text), quote_style))
        elif block.kind == "table" and block.rows:
            column_count = len(block.rows[0])
            available_width = page_size[0] - 1.2 * inch
            weights = []
            for index in range(column_count):
                longest = max(len(str(row[index])) for row in block.rows)
                weights.append(max(4, min(longest, 24)))
            total_weight = sum(weights)
            widths = [available_width * weight / total_weight for weight in weights]
            cell_style = ParagraphStyle(
                "ChineseTableCell",
                parent=body,
                fontSize=6.5 if column_count > 6 else 8,
                leading=9 if column_count > 6 else 11,
                spaceAfter=0,
            )
            table_data = [
                [Paragraph(html.escape(cell), cell_style) for cell in row]
                for row in block.rows
            ]
            table = Table(table_data, colWidths=widths, repeatRows=1, hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CBD5E1")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.extend([table, Spacer(1, 8)])
        else:
            story.append(Paragraph(_pdf_inline(block.text), body))

    def footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("STSong-Light", 8)
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.drawCentredString(page_size[0] / 2, 0.35 * inch, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return stream.getvalue()


def _xlsx_value(value: object) -> object:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return json.dumps(value, ensure_ascii=False, default=str)


def _safe_sheet_name(value: str, used: set[str]) -> str:
    base = re.sub(r"[\\/*?:\[\]]", "-", value).strip()[:31] or "数据"
    candidate = base
    suffix = 2
    while candidate in used:
        marker = f"-{suffix}"
        candidate = f"{base[:31 - len(marker)]}{marker}"
        suffix += 1
    used.add(candidate)
    return candidate


def _xlsx_payload(
    title: str,
    markdown: str,
    data_sets: list[tuple[str, list[str], list[dict[str, object]]]],
) -> bytes:
    """生成带报告摘要和独立数据工作表的 Excel 文件。"""

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    workbook = Workbook()
    report_sheet = workbook.active
    report_sheet.title = "报告"
    report_sheet.sheet_view.showGridLines = False
    report_sheet.merge_cells("A1:F1")
    title_cell = report_sheet["A1"]
    title_cell.value = _plain_text(title)
    title_cell.font = Font(name="微软雅黑", size=18, bold=True, color="0B2545")
    title_cell.alignment = Alignment(vertical="center")
    report_sheet.row_dimensions[1].height = 32
    row_index = 3
    for block in _parse_markdown(markdown):
        if block.kind == "table" and block.rows:
            for table_row_index, source_row in enumerate(block.rows):
                for column_index, value in enumerate(source_row, start=1):
                    cell = report_sheet.cell(row=row_index, column=column_index, value=value)
                    cell.alignment = Alignment(vertical="top", wrap_text=True)
                    if table_row_index == 0:
                        cell.font = Font(name="微软雅黑", bold=True, color="FFFFFF")
                        cell.fill = PatternFill("solid", fgColor="1F4D78")
                row_index += 1
            row_index += 1
            continue
        report_sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=6)
        cell = report_sheet.cell(row=row_index, column=1, value=block.text)
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        if block.kind == "heading":
            cell.font = Font(
                name="微软雅黑",
                size={1: 15, 2: 13, 3: 11}.get(block.level, 11),
                bold=True,
                color="2E74B5",
            )
            report_sheet.row_dimensions[row_index].height = 24
        else:
            prefix = "• " if block.kind == "bullet" else ""
            if prefix:
                cell.value = f"{prefix}{block.text}"
            cell.font = Font(name="微软雅黑", size=10, color="202124")
            report_sheet.row_dimensions[row_index].height = 30
        row_index += 1
    for column_index in range(1, 7):
        report_sheet.column_dimensions[get_column_letter(column_index)].width = 18

    used_names = {"报告"}
    thin = Side(style="thin", color="D9E0E8")
    for data_name, columns, rows in data_sets:
        sheet = workbook.create_sheet(_safe_sheet_name(data_name, used_names))
        sheet.sheet_view.showGridLines = False
        sheet.freeze_panes = "A2"
        for column_index, column in enumerate(columns, start=1):
            cell = sheet.cell(row=1, column=column_index, value=column)
            cell.font = Font(name="微软雅黑", bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4D78")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row_index, row in enumerate(rows, start=2):
            for column_index, column in enumerate(columns, start=1):
                cell = sheet.cell(
                    row=row_index,
                    column=column_index,
                    value=_xlsx_value(row.get(column)),
                )
                cell.font = Font(name="微软雅黑", size=10)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(bottom=thin)
        if columns:
            sheet.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{max(1, len(rows) + 1)}"
        for column_index, column in enumerate(columns, start=1):
            values = [len(str(row.get(column) or "")) for row in rows[:200]]
            width = max([len(column) * 1.8, *values, 10])
            sheet.column_dimensions[get_column_letter(column_index)].width = min(width + 2, 40)
        sheet.row_dimensions[1].height = 28

    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def create_report_files(
    task_id: int,
    title: str,
    markdown: str,
    observations: list[dict[str, Any]],
    settings: Settings | None = None,
) -> list[dict[str, object]]:
    """把最终报告及其数据证据保存到现有 MinIO 的任务目录。"""

    current = settings or get_settings()
    if not current.report_files_enabled:
        return []
    artifacts: list[dict[str, object]] = []
    prefix = f"tasks/{task_id}"
    markdown_body = f"# {title}\n\n{markdown.strip()}\n"
    data_sets = _data_observations(observations)

    report_files = [
        (
            "分析报告.md",
            "report",
            f"{prefix}/report.md",
            markdown_body.encode("utf-8"),
            "text/markdown; charset=utf-8",
        ),
        (
            "分析报告.docx",
            "report",
            f"{prefix}/report.docx",
            _docx_payload(title, markdown),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        (
            "分析报告.pdf",
            "report",
            f"{prefix}/report.pdf",
            _pdf_payload(title, markdown),
            "application/pdf",
        ),
        (
            "数据报告.xlsx",
            "query_workbook",
            f"{prefix}/report.xlsx",
            _xlsx_payload(title, markdown, data_sets),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    ]
    for name, kind, object_key, payload, content_type in report_files:
        stored = save_file(object_key, payload, content_type)
        artifacts.append({"name": name, "kind": kind, **stored})

    evidence_body = json.dumps(
        {"taskId": task_id, "observations": observations},
        ensure_ascii=False,
        indent=2,
        default=str,
    ).encode("utf-8")
    evidence = save_file(
        f"{prefix}/evidence.json",
        evidence_body,
        "application/json",
    )
    artifacts.append({"name": "证据清单.json", "kind": "evidence", **evidence})

    for data_name, columns, rows in data_sets:
        csv_file = save_file(
            f"{prefix}/{data_name}.csv",
            _csv_payload(rows, columns),
            "text/csv; charset=utf-8",
        )
        artifacts.append(
            {"name": f"{data_name}.csv", "kind": "query_data", **csv_file}
        )
    logger.info("Agent report files created: task_id=%s files=%d", task_id, len(artifacts))
    return artifacts
