"""企业自建 RAG 知识库。

PostgreSQL 负责文档台账，MinIO 保存原文件，Qdrant 保存分块向量。本模块只处理
文件解析、向量入库和带权限过滤的检索；任何浏览器请求都必须先经过 NestJS RBAC。
"""

import hashlib
import io
import json
import logging
import re
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

import httpx
from docx import Document
from qdrant_client import models

from app.config.settings import Settings, get_settings
from app.integrations import minio_client, qdrant_client, save_file

logger = logging.getLogger(__name__)


class KnowledgeBaseError(RuntimeError):
    """可以安全返回给内部调用方的知识库错误。"""


@dataclass(frozen=True)
class KnowledgeConfig:
    """无需重启即可调整的检索与分块参数。"""

    enabled: bool = True
    collection_name: str = "enterprise_knowledge"
    chunk_size: int = 900
    chunk_overlap: int = 150
    top_k: int = 5
    candidate_multiplier: int = 3
    score_threshold: float = 0.25
    semantic_weight: float = 0.85
    max_context_chars: int = 14000
    max_file_bytes: int = 20 * 1024 * 1024
    embedding_batch_size: int = 16


@dataclass(frozen=True)
class KnowledgeResult:
    """可直接注入大模型上下文的知识库检索结果。"""

    enabled: bool
    context: str
    hit_count: int


def load_knowledge_config(settings: Settings | None = None) -> KnowledgeConfig:
    """每次操作都读取 JSON，修改服务器文件后无需重启服务。"""

    current = settings or get_settings()
    path = Path(current.knowledge_config_file)
    if not path.is_file():
        raise KnowledgeBaseError(f"知识库配置文件不存在：{path}")
    try:
        raw = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        raise KnowledgeBaseError("知识库配置文件无法读取或 JSON 格式错误") from error
    if not isinstance(raw, dict):
        raise KnowledgeBaseError("知识库配置文件顶层必须是 JSON 对象")
    try:
        config = KnowledgeConfig(
            enabled=bool(raw.get("enabled", True)),
            collection_name=str(
                raw.get("collection_name", "enterprise_knowledge")
            ).strip(),
            chunk_size=int(raw.get("chunk_size", 900)),
            chunk_overlap=int(raw.get("chunk_overlap", 150)),
            top_k=int(raw.get("top_k", 5)),
            candidate_multiplier=int(raw.get("candidate_multiplier", 3)),
            score_threshold=float(raw.get("score_threshold", 0.25)),
            semantic_weight=float(raw.get("semantic_weight", 0.85)),
            max_context_chars=int(raw.get("max_context_chars", 14000)),
            max_file_bytes=int(raw.get("max_file_bytes", 20 * 1024 * 1024)),
            embedding_batch_size=int(raw.get("embedding_batch_size", 16)),
        )
    except (TypeError, ValueError) as error:
        raise KnowledgeBaseError("知识库配置文件包含无效参数类型") from error
    if not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_-]{1,62}", config.collection_name):
        raise KnowledgeBaseError("collection_name 格式无效")
    if not 200 <= config.chunk_size <= 4000:
        raise KnowledgeBaseError("chunk_size 必须在 200 到 4000 之间")
    if not 0 <= config.chunk_overlap < config.chunk_size:
        raise KnowledgeBaseError("chunk_overlap 必须小于 chunk_size")
    if not 1 <= config.top_k <= 20 or not 1 <= config.candidate_multiplier <= 10:
        raise KnowledgeBaseError("top_k 或 candidate_multiplier 超出允许范围")
    if not 0 <= config.score_threshold <= 1 or not 0 <= config.semantic_weight <= 1:
        raise KnowledgeBaseError("检索阈值和权重必须在 0 到 1 之间")
    if not 1000 <= config.max_context_chars <= 100000:
        raise KnowledgeBaseError("max_context_chars 超出允许范围")
    if not 1024 <= config.max_file_bytes <= 100 * 1024 * 1024:
        raise KnowledgeBaseError("max_file_bytes 超出允许范围")
    if not 1 <= config.embedding_batch_size <= 128:
        raise KnowledgeBaseError("embedding_batch_size 超出允许范围")
    return config


def _extract_text(file_name: str, content: bytes) -> str:
    """从受支持文件中提取纯文本；不执行文件内宏或脚本。"""

    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json"}:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise KnowledgeBaseError("文本文件必须使用 UTF-8 编码") from error
    if suffix == ".docx":
        try:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                paragraphs.extend(
                    " | ".join(cell.text for cell in row.cells) for row in table.rows
                )
            return "\n".join(paragraphs)
        except Exception as error:
            raise KnowledgeBaseError("Word 文件损坏或无法解析") from error
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            return "\n".join(
                page.extract_text() or ""
                for page in PdfReader(io.BytesIO(content)).pages
            )
        except Exception as error:
            raise KnowledgeBaseError("PDF 文件损坏、加密或无法解析") from error
    raise KnowledgeBaseError("仅支持 PDF、DOCX、TXT、Markdown、CSV 和 JSON 文件")


def _chunks(text: str, size: int, overlap: int) -> list[str]:
    """按段落优先切块，超长段落再按字符窗口切分。"""

    normalized = re.sub(r"\r\n?", "\n", text)
    normalized = re.sub(r"[ \t]+", " ", normalized).strip()
    if not normalized:
        raise KnowledgeBaseError("文件没有可索引的文本内容")
    pieces = [
        piece.strip() for piece in re.split(r"\n{2,}", normalized) if piece.strip()
    ]
    output: list[str] = []
    current = ""
    for piece in pieces:
        if len(piece) > size:
            if current:
                output.append(current)
                current = ""
            step = size - overlap
            output.extend(
                piece[start : start + size] for start in range(0, len(piece), step)
            )
        elif not current:
            current = piece
        elif len(current) + 2 + len(piece) <= size:
            current = f"{current}\n\n{piece}"
        else:
            output.append(current)
            prefix = current[-overlap:] if overlap else ""
            current = f"{prefix}\n\n{piece}".strip()
    if current:
        output.append(current)
    return [chunk for chunk in output if chunk.strip()]


async def _embeddings(
    texts: list[str], settings: Settings, batch_size: int
) -> list[list[float]]:
    """通过 OpenAI 兼容接口批量生成真实语义向量。"""

    api_key = (
        settings.embedding_api_key.get_secret_value()
        if settings.embedding_api_key
        else ""
    )
    if not api_key:
        raise KnowledgeBaseError("EMBEDDING_API_KEY 未配置")
    vectors: list[list[float]] = []
    timeout = httpx.Timeout(settings.embedding_request_timeout_seconds, connect=10)
    try:
        async with httpx.AsyncClient(timeout=timeout) as client:
            for start in range(0, len(texts), batch_size):
                response = await client.post(
                    f"{settings.embedding_base_url.rstrip('/')}/embeddings",
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={
                        "model": settings.embedding_model,
                        "input": texts[start : start + batch_size],
                    },
                )
                if response.status_code >= 400:
                    raise KnowledgeBaseError(
                        f"向量模型请求失败（HTTP {response.status_code}）"
                    )
                body = response.json()
                data = body.get("data") if isinstance(body, dict) else None
                if not isinstance(data, list):
                    raise KnowledgeBaseError("向量模型返回了无效数据")
                ordered = sorted(data, key=lambda item: int(item.get("index", 0)))
                vectors.extend(
                    [list(map(float, item["embedding"])) for item in ordered]
                )
    except KnowledgeBaseError:
        raise
    except (httpx.HTTPError, ValueError, KeyError, TypeError) as error:
        raise KnowledgeBaseError("无法调用向量模型") from error
    if len(vectors) != len(texts) or any(
        len(vector) != settings.embedding_dimension for vector in vectors
    ):
        raise KnowledgeBaseError("向量模型返回数量或维度与 EMBEDDING_DIMENSION 不一致")
    return vectors


def _ensure_collection(config: KnowledgeConfig, settings: Settings) -> None:
    """创建独立知识库集合及权限过滤索引。"""

    client = qdrant_client()
    if not client.collection_exists(config.collection_name):
        client.create_collection(
            collection_name=config.collection_name,
            vectors_config=models.VectorParams(
                size=settings.embedding_dimension, distance=models.Distance.COSINE
            ),
        )
        client.create_payload_index(
            config.collection_name,
            "visibility",
            models.PayloadSchemaType.KEYWORD,
            wait=True,
        )
        client.create_payload_index(
            config.collection_name,
            "document_id",
            models.PayloadSchemaType.INTEGER,
            wait=True,
        )


async def ingest_document(
    document_id: int,
    title: str,
    file_name: str,
    content_type: str,
    visibility: Literal["public", "admin"],
    content: bytes,
) -> dict[str, object]:
    """保存原文件、分块、生成向量并写入 Qdrant。"""

    config = load_knowledge_config()
    settings = get_settings()
    if not config.enabled:
        raise KnowledgeBaseError("知识库当前已停用")
    if len(content) > config.max_file_bytes:
        raise KnowledgeBaseError("文件超过知识库配置允许的最大大小")
    chunks = _chunks(
        _extract_text(file_name, content), config.chunk_size, config.chunk_overlap
    )
    vectors = await _embeddings(chunks, settings, config.embedding_batch_size)
    _ensure_collection(config, settings)
    checksum = hashlib.sha256(content).hexdigest()
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", Path(file_name).name) or "source"
    object_key = f"knowledge/{document_id}/source/{safe_name}"
    save_file(object_key, content, content_type or "application/octet-stream")
    points = []
    for index, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True)):
        point_id = str(
            uuid.uuid5(
                uuid.NAMESPACE_URL, f"enterprise-knowledge:{document_id}:{index}"
            )
        )
        points.append(
            models.PointStruct(
                id=point_id,
                vector=vector,
                payload={
                    "document_id": document_id,
                    "title": title,
                    "file_name": file_name,
                    "visibility": visibility,
                    "chunk_index": index,
                    "text": chunk,
                    "checksum": checksum,
                },
            )
        )
    try:
        qdrant_client().upsert(
            collection_name=config.collection_name, points=points, wait=True
        )
    except Exception:
        # 向量写入失败时补偿删除原文件，避免索引失败产生不可管理的孤儿对象。
        try:
            storage = minio_client()
            if storage.bucket_exists(settings.minio_bucket):
                storage.remove_object(settings.minio_bucket, object_key)
        except Exception:
            logger.exception(
                "knowledge object rollback failed: document_id=%d", document_id
            )
        raise
    logger.info(
        "knowledge document indexed: document_id=%d chunks=%d", document_id, len(points)
    )
    return {"objectKey": object_key, "chunkCount": len(points), "checksum": checksum}


def delete_document(document_id: int, object_key: str | None = None) -> None:
    """幂等删除文档的全部向量和 MinIO 原文件。"""

    config = load_knowledge_config()
    client = qdrant_client()
    if client.collection_exists(config.collection_name):
        client.delete(
            collection_name=config.collection_name,
            points_selector=models.FilterSelector(
                filter=models.Filter(
                    must=[
                        models.FieldCondition(
                            key="document_id",
                            match=models.MatchValue(value=document_id),
                        )
                    ]
                )
            ),
            wait=True,
        )
    storage = minio_client()
    settings = get_settings()
    if object_key and storage.bucket_exists(settings.minio_bucket):
        storage.remove_object(settings.minio_bucket, object_key)
    logger.info("knowledge document deleted: document_id=%d", document_id)


def update_document_visibility(
    document_id: int, visibility: Literal["public", "admin"]
) -> None:
    """批量修改全部分块的可见性 payload。"""

    config = load_knowledge_config()
    client = qdrant_client()
    if client.collection_exists(config.collection_name):
        client.set_payload(
            collection_name=config.collection_name,
            payload={"visibility": visibility},
            points=models.Filter(
                must=[
                    models.FieldCondition(
                        key="document_id", match=models.MatchValue(value=document_id)
                    )
                ]
            ),
            wait=True,
        )


def _terms(text: str) -> set[str]:
    """生成轻量词项，用于语义召回后的关键词重排。"""

    lowered = text.lower()
    words = set(re.findall(r"[a-z0-9_]{2,}|[\u4e00-\u9fff]{2,}", lowered))
    chinese = "".join(re.findall(r"[\u4e00-\u9fff]", lowered))
    words.update(
        chinese[index : index + 2] for index in range(max(0, len(chinese) - 1))
    )
    return words


async def retrieve_knowledge(query: str, allow_admin: bool = False) -> KnowledgeResult:
    """语义召回后重排，并在 Qdrant 查询阶段执行文档 ACL。"""

    config = load_knowledge_config()
    if not config.enabled:
        return KnowledgeResult(enabled=False, context="", hit_count=0)
    settings = get_settings()
    client = qdrant_client()
    if not client.collection_exists(config.collection_name):
        return KnowledgeResult(enabled=True, context="", hit_count=0)
    query_vector = (await _embeddings([query.strip()[:1000]], settings, 1))[0]
    acl_filter = models.Filter(
        must=[
            models.FieldCondition(
                key="visibility",
                match=(
                    models.MatchAny(any=["public", "admin"])
                    if allow_admin
                    else models.MatchValue(value="public")
                ),
            )
        ]
    )
    result = client.query_points(
        collection_name=config.collection_name,
        query=query_vector,
        query_filter=acl_filter,
        limit=config.top_k * config.candidate_multiplier,
        with_payload=True,
    )
    query_terms = _terms(query)
    ranked: list[tuple[float, object]] = []
    for point in result.points:
        payload = point.payload or {}
        text = str(payload.get("text") or "")
        semantic = max(0.0, min(float(point.score), 1.0))
        terms = _terms(text)
        lexical = len(query_terms & terms) / max(1, len(query_terms))
        score = (
            config.semantic_weight * semantic + (1 - config.semantic_weight) * lexical
        )
        if score >= config.score_threshold:
            ranked.append((score, point))
    ranked.sort(key=lambda item: item[0], reverse=True)
    blocks: list[str] = []
    used = 0
    for score, point in ranked[: config.top_k]:
        payload = point.payload or {}
        header = f"[知识库 {len(blocks) + 1}｜来源：{payload.get('title') or payload.get('file_name') or '未知文档'}｜相关度：{score:.3f}]"
        text = str(payload.get("text") or "").strip()
        remaining = config.max_context_chars - used - len(header) - 1
        if remaining <= 0:
            break
        block = f"{header}\n{text[:remaining]}"
        blocks.append(block)
        used += len(block)
    logger.info(
        "knowledge retrieval completed: admin_scope=%s hits=%d",
        allow_admin,
        len(blocks),
    )
    return KnowledgeResult(
        enabled=True, context="\n\n".join(blocks), hit_count=len(blocks)
    )
